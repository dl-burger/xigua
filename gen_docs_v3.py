import os, re, html
from docx import Document
from docx.oxml.ns import qn

base = r"C:/Users/m1594/WorkBuddy/2026-06-25-04-43-16/outputs/portfolio"

files = {
    "elevator_gdd_master": "游戏设计总纲",
    "elevator_pm_master": "项目管理文档",
    "elevator_art_master": "美术需求文档",
    "elevator_audio_master": "音效需求文档",
}

CSS = """<style>
  :root {
    --bg: #f9f7f4; --card: #ffffff; --text: #1e1b18;
    --text-secondary: #5c554e; --text-muted: #9b9489; --accent: #b8753e;
    --accent-dim: #d4a76a; --accent-subtle: #f5ede0; --border: #e8e2d8;
    --highlight-bg: #faf7f0; --shadow-sm: 0 1px 3px rgba(80,60,30,0.04);
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", sans-serif;
    background: var(--bg); color: var(--text); line-height: 1.85; -webkit-font-smoothing: antialiased;
  }
  body::before {
    content: ""; position: fixed; top: 0; left: 0; right: 0; bottom: 0;
    background-image: url("data:image/svg+xml,%3Csvg width='60' height='60' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M30 5 Q35 15 30 25 Q25 15 30 5Z' fill='%23d4c8b0' fill-opacity='0.04'/%3E%3C/svg%3E");
    background-size: 60px 60px; pointer-events: none; z-index: -1;
  }
  .wrapper { max-width: 860px; margin: 0 auto; padding: 3rem 2rem; }
  .back { display: inline-block; margin-bottom: 2rem; color: var(--accent); text-decoration: none; font-size: 0.85rem; }
  .back:hover { text-decoration: underline; }
  .doc-title { font-size: 1.8rem; font-weight: 700; color: var(--text); margin-bottom: 0.5rem; }
  .doc-sub { font-size: 1rem; color: var(--text-secondary); margin-bottom: 2rem; }
  .section { margin-bottom: 2.5rem; }
  .section h2 {
    font-size: 1.25rem; font-weight: 600; margin-bottom: 1rem; color: var(--accent);
    display: flex; align-items: center; gap: 0.6rem;
  }
  .section h2::before { content: ""; display: inline-block; width: 3px; height: 1.2em; background: var(--accent); border-radius: 2px; }
  .section h3 { font-size: 1.05rem; font-weight: 600; margin: 1.5rem 0 0.75rem; color: var(--text); }
  .section h4 { font-size: 0.95rem; font-weight: 600; margin: 1rem 0 0.5rem; color: var(--text-secondary); }
  .section p { color: var(--text-secondary); font-size: 0.92rem; margin-bottom: 0.75rem; line-height: 1.85; }
  .table-wrap { overflow-x: auto; margin: 1rem 0 1.5rem; }
  table {
    width: 100%; border-collapse: collapse; font-size: 0.82rem;
    background: var(--card); border-radius: 8px; overflow: hidden; border: 1px solid var(--border);
  }
  th { background: var(--accent-subtle); padding: 0.5rem 0.7rem; text-align: left; font-weight: 600; color: var(--accent); font-size: 0.76rem; letter-spacing: 0.04em; white-space: nowrap; }
  td { padding: 0.45rem 0.7rem; border-top: 1px solid var(--border); color: var(--text-secondary); vertical-align: top; }
  tr:nth-child(even) td { background: var(--highlight-bg); }
  footer { text-align: center; color: var(--text-muted); font-size: 0.8rem; padding: 2rem; }
  @media (max-width: 768px) {
    .wrapper { padding: 2rem 1.25rem; }
    .doc-title { font-size: 1.4rem; }
  }
</style>"""

def render_docx_table(table):
    """Render a python-docx Table object to HTML, handling merged cells properly"""
    # Build grid with rowspan/colspan info
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = cell.text.strip()
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            grid_span = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    grid_span = int(gs.get(qn('w:val'), '1'))
            cells.append((text, grid_span))
        rows.append(cells)
    
    if not rows:
        return ""
    
    # Max columns
    max_cols = max(sum(c[1] for c in row) for row in rows)
    
    # First row check for header
    is_header = len(rows) > 1
    
    h = '<div class="table-wrap"><table>'
    
    start = 0
    if is_header:
        h += '<thead><tr>'
        for text, span in rows[0]:
            attrs = f' colspan="{span}"' if span > 1 else ''
            h += f'<th{attrs}>{html.escape(text)}</th>'
        h += '</tr></thead><tbody>'
        start = 1
    else:
        h += '<tbody>'
    
    for row in rows[start:]:
        h += '<tr>'
        for text, span in row:
            attrs = f' colspan="{span}"' if span > 1 else ''
            h += f'<td{attrs}>{html.escape(text)}</td>'
        h += '</tr>'
    
    h += '</tbody></table></div>'
    return h

def docx_to_html(docx_path):
    """Convert docx to HTML body, interleaving paragraphs and tables"""
    doc = Document(docx_path)
    body_parts = []
    
    # Collect all body elements in order
    elements = []
    for child in doc.element.body:
        if child.tag == qn('w:p'):
            # It's a paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para:
                elements.append(('p', para))
        elif child.tag == qn('w:tbl'):
            # It's a table
            tbl = None
            for t in doc.tables:
                if t._element is child:
                    tbl = t
                    break
            if tbl:
                elements.append(('table', tbl))
    
    current_section_title = None
    section_content = []
    sections = []
    
    for etype, obj in elements:
        if etype == 'p':
            text = obj.text.strip()
            style = obj.style.name if obj.style else ''
            
            if not text:
                continue
            
            # Detect headings
            if style.startswith('Heading') or text.startswith('### '):
                if style.startswith('Heading'):
                    level = int(style.replace('Heading ', '').split()[0]) if 'Heading' in style else 1
                    title = text
                else:
                    title = text[4:]
                    level = 1
                
                # Flush previous section
                if current_section_title:
                    sections.append((current_section_title, '\n'.join(section_content)))
                    section_content = []
                current_section_title = (title, level)
            else:
                section_content.append(text)
        
        elif etype == 'table':
            table_html = render_docx_table(obj)
            section_content.append(f'<!--TABLE-->{table_html}')
    
    # Flush last section
    if current_section_title:
        sections.append((current_section_title, '\n'.join(section_content)))
    
    # Render sections to HTML
    html_parts = []
    for (title, level), content in sections:
        tag = 'h2' if level <= 2 else 'h3'
        html_parts.append(f'<div class="section"><{tag}>{html.escape(title)}</{tag}>')
        
        # Split content by TABLE markers
        content_parts = re.split(r'<!--TABLE-->', content)
        for part in content_parts:
            if part.startswith('<div class="table-wrap">'):
                html_parts.append(part)
            elif part.strip():
                # Regular text
                lines = part.strip().split('\n')
                para = []
                for line in lines:
                    line = line.strip()
                    if not line:
                        if para:
                            html_parts.append(f'<p>{" ".join(html.escape(p) for p in para)}</p>')
                            para = []
                        continue
                    if line.startswith(('●','•','-','·')) or re.match(r'^\d+[\.\)]', line):
                        if para:
                            html_parts.append(f'<p>{" ".join(html.escape(p) for p in para)}</p>')
                            para = []
                        html_parts.append(f'<p style="padding-left:1rem;">{html.escape(line)}</p>')
                    elif re.match(r'^[一二三四五六七八九十]、', line):
                        if para:
                            html_parts.append(f'<p>{" ".join(html.escape(p) for p in para)}</p>')
                            para = []
                        html_parts.append(f'<h3>{html.escape(line)}</h3>')
                    elif re.match(r'^\d+\.\d+\s', line) or line.startswith('★'):
                        if para:
                            html_parts.append(f'<p>{" ".join(html.escape(p) for p in para)}</p>')
                            para = []
                        html_parts.append(f'<h4>{html.escape(line)}</h4>')
                    else:
                        para.append(line)
                if para:
                    html_parts.append(f'<p>{" ".join(html.escape(p) for p in para)}</p>')
        
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)


for fname, title in files.items():
    docx_path = os.path.join(base, f"{fname}.docx")
    if not os.path.exists(docx_path):
        print(f"SKIP: {docx_path}")
        continue
    
    body = docx_to_html(docx_path)
    
    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} — Elevator — Don't Touch</title>
{CSS}
</head>
<body>
<div class="wrapper">
  <a href="index.html" class="back">← 返回作品集</a>
  <h1 class="doc-title">{title}</h1>
  <p class="doc-sub">Elevator — Don't Touch（别按那个键）· 第一人称 3D 密室互动恐怖游戏</p>
  {body}
</div>
<footer><p>© 2026 代立强 · Elevator — Don't Touch 项目文档</p></footer>
</body>
</html>"""
    
    out_path = os.path.join(base, f"{fname}.html")
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    # Count tables
    import re
    table_count = len(re.findall(r'<div class="table-wrap">', html_content))
    print(f"Generated: {fname}.html ({table_count} tables)")

print("Done!")
